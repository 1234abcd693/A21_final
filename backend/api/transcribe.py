"""
语音转文字 + 文档解析 + 确认导入 + 报告导出
"""

import os
import json
import tempfile
import uuid
from typing import Optional

from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from tools.parser import parse_file
from tools.extractor import extract_knowledge
from core.config import settings

router = APIRouter()

# ---- 对话消息缓存（feedback 需要 question/answer 上下文） ----
# 实际生产应存数据库，这里用内存缓存
_message_cache: dict[str, dict] = {}


def cache_message(message_id: str, question: str, answer: str, chunks: str):
    _message_cache[message_id] = {
        "question": question, "answer_text": answer, "retrieved_chunks": chunks,
    }


# ===== 语音转文字 =====

@router.post("/transcribe")
async def transcribe(audio: UploadFile = File(...)):
    """语音转文字 → whisper-server"""
    tmp_path = os.path.join(tempfile.gettempdir(), f"audio_{uuid.uuid4().hex}.wav")
    with open(tmp_path, "wb") as f:
        f.write(await audio.read())

    import httpx
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            with open(tmp_path, "rb") as f:
                resp = await client.post(
                    f"{settings.WHISPER_SERVER_URL}/inference",
                    files={"file": ("audio.wav", f, "audio/wav")},
                )
            result = resp.json() if resp.status_code == 200 else {"text": ""}
    except Exception:
        result = {"text": ""}
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return {"text": result.get("text", ""), "duration_ms": 0, "processing_ms": 0}


# ===== 文档解析与知识抽取 =====

@router.post("/parse")
async def parse_document(file: UploadFile = File(...)):
    """上传文档 → 解析 → 知识抽取"""
    tmp_path = os.path.join(tempfile.gettempdir(), file.filename or "upload")
    with open(tmp_path, "wb") as f:
        f.write(await file.read())

    try:
        parsed = parse_file(tmp_path)
    except ValueError as e:
        raise HTTPException(400, str(e))
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    extracted = extract_knowledge(parsed["text"])
    parse_id = f"parse_{uuid.uuid4().hex[:8]}"

    return {
        "parse_id": parse_id,
        "file_name": file.filename,
        "extracted_text_length": parsed["text_length"],
        "candidates": extracted["candidates"],
        "regex_count": extracted["regex_count"],
        "model_count": extracted["model_count"],
    }


# ===== 确认导入 =====

class ConfirmRequest(BaseModel):
    parse_id: str
    confirmed: list[dict] = []
    rejected: list[str] = []
    edits: list[dict] = []


@router.post("/confirm")
async def confirm_import(req: ConfirmRequest):
    """确认导入知识抽取结果到 Neo4j + Chroma"""
    if not req.confirmed:
        return {"status": "ok", "imported": {"entities": 0, "relationships": 0}, "chroma_chunks_added": 0}

    entities_added = 0
    relationships_added = 0
    chroma_chunks = []

    try:
        from kg.neo4j_client import _get_driver
        with _get_driver().session() as session:
            for cand in req.confirmed:
                entity_type = cand.get("type", "")
                name = cand.get("name", "")
                uid = cand.get("uid", f"IMP_{uuid.uuid4().hex[:6]}")
                source = cand.get("source_doc", req.parse_id)
                page = cand.get("source_page", "")

                # MERGE 幂等创建实体
                props = {k: v for k, v in cand.items() if k not in ("uid", "type", "relations", "source", "confidence")}
                props["uid"] = uid
                props.setdefault("name", name)
                props.setdefault("source_doc", source)
                props.setdefault("source_page", page)

                session.run(
                    f"MERGE (n:{entity_type} {{uid: $uid}}) SET n += $props",
                    uid=uid, props=props,
                )
                entities_added += 1

                # 创建关系
                for rel in cand.get("relations", []):
                    target_uid = rel.get("target", "")
                    rel_type = rel.get("type", "")
                    if target_uid and rel_type:
                        try:
                            session.run(
                                f"MATCH (a {{uid: $from}}), (b {{uid: $to}}) MERGE (a)-[:{rel_type}]->(b)",
                                **{"from": uid, "to": target_uid},
                            )
                            relationships_added += 1
                        except Exception:
                            pass

                # 收集 Chroma 数据
                description = (
                    cand.get("description", "")
                    or f"{entity_type}: {name}"
                )
                chroma_chunks.append({
                    "id": f"chunk_{uid}",
                    "text": f"{entity_type}: {name}\n{description}",
                    "metadata": {"doc_name": source, "page": page, "doc_type": "uploaded", "entity_type": entity_type},
                })
    except Exception as e:
        raise HTTPException(500, f"导入失败: {str(e)}")

    # 写入 Chroma
    chroma_added = 0
    if chroma_chunks:
        try:
            from rag.vector_store import add_chunks
            add_chunks(
                ids=[c["id"] for c in chroma_chunks],
                documents=[c["text"] for c in chroma_chunks],
                metadatas=[c["metadata"] for c in chroma_chunks],
            )
            chroma_added = len(chroma_chunks)
        except Exception:
            pass

    return {
        "status": "ok",
        "imported": {"entities": entities_added, "relationships": relationships_added},
        "chroma_chunks_added": chroma_added,
    }


# ===== 报告导出 =====

@router.get("/report")
async def export_report(session_id: Optional[str] = None):
    """导出问答报告（纯文本，Word 需要 python-docx）"""
    messages = []
    if session_id:
        from data.database import get_connection
        conn = get_connection()
        rows = conn.execute(
            "SELECT role, content, created_at FROM messages WHERE session_id = ? ORDER BY created_at",
            [session_id],
        ).fetchall()
        conn.close()
        messages = [dict(r) for r in rows]

    # 生成纯文本报告
    lines = ["A21 船舶故障诊断系统 - 问答报告", "=" * 40, ""]
    for msg in messages:
        role = "用户" if msg["role"] == "user" else "系统"
        lines.append(f"[{role}] {msg['created_at']}")
        lines.append(msg["content"])
        lines.append("")

    report_text = "\n".join(lines)

    # 尝试生成 Word
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("A21 船舶故障诊断系统 - 问答报告", 0)
        for msg in messages:
            role = "用户" if msg["role"] == "user" else "系统"
            doc.add_paragraph(f"[{role}] {msg['created_at']}", style="List Bullet")
            doc.add_paragraph(msg["content"])
            doc.add_paragraph("")

        tmp_path = os.path.join(tempfile.gettempdir(), f"report_{uuid.uuid4().hex[:8]}.docx")
        doc.save(tmp_path)
        return FileResponse(tmp_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="问答报告.docx")
    except ImportError:
        pass

    return {"report": report_text, "message_count": len(messages)}
